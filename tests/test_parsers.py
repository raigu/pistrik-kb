import pytest
from pathlib import Path

FIXTURES = Path(__file__).parent / "fixtures"


def test_parse_markdown():
    from parsers import parse_markdown

    md_file = FIXTURES / "sample.md"
    result = parse_markdown(md_file)
    assert "Sample heading" in result
    assert isinstance(result, str)


def test_parse_html():
    from parsers import parse_html

    html = "<h1>Title</h1><p>Body text</p><script>var x=1;</script>"
    result = parse_html(html)
    assert "Title" in result
    assert "Body text" in result
    assert "var x" not in result  # scripts stripped


def test_parse_pdf(tmp_path):
    from parsers import parse_pdf
    import pymupdf

    # Create a minimal PDF with text
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "PISTRIK test document")
    pdf_path = tmp_path / "test.pdf"
    doc.save(str(pdf_path))
    doc.close()

    result = parse_pdf(pdf_path)
    assert "PISTRIK test document" in result


def test_parse_docx(tmp_path):
    from parsers import parse_docx
    from docx import Document

    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("Test paragraph content.")
    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))

    result = parse_docx(docx_path)
    assert "Test Heading" in result
    assert "Test paragraph content" in result


def test_parse_text(tmp_path):
    from parsers import parse_text

    txt_path = tmp_path / "note.txt"
    txt_path.write_text("Meeting notes: PISTRIK API changed")

    result = parse_text(txt_path)
    assert "PISTRIK API changed" in result


def test_parse_file_dispatches_by_extension(tmp_path):
    from parsers import parse_file

    txt_path = tmp_path / "note.md"
    txt_path.write_text("# Hello\nWorld")

    result = parse_file(txt_path)
    assert "Hello" in result


def test_parse_file_unsupported_extension(tmp_path):
    from parsers import parse_file

    path = tmp_path / "data.xlsx"
    path.write_bytes(b"fake")

    with pytest.raises(ValueError, match="Unsupported"):
        parse_file(path)
